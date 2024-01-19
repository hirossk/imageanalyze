import PySimpleGUI as sg
#wordのファイルを操作するため
import docx
#AWSのサービスを利用するため
import boto3
#wordを起動するため
import win32com.client
#docxをpdfに変換するため
from docx2pdf import convert
#Adobe Reader（外部プログラム）を動かすため
import subprocess
#ディレクトリフォルダーを取得するため
import os

#カレントディレクトリを取得する
cwd = os.getcwd() + "\\"

#東京リージョン
REGION = 'ap-northeast-1'
SRC_LANG = 'en'
TRG_LANG = 'ja'

#翻訳エンジンへの接続
translate = boto3.client('translate', region_name=REGION)

#初期バージョン（翻訳のみ）
layout = [
   [sg.Text("英語で書かれているwordの文書を選択してください")],
   [sg.Text("ファイル"), sg.InputText(), sg.FileBrowse(key="file1",button_text="ファイル選択")],
   [sg.Submit("翻訳"), sg.Submit("保存"), sg.Cancel("終了")],
]
#保存バージョン（翻訳・保存）
#layout = [
#   [sg.Text("英語で書かれているwordの文書を選択してください")],
#   [sg.Text("ファイル"), sg.InputText(), sg.FileBrowse(key="file1",button_text="ファイル選択")],
#   [sg.Submit("翻訳"), sg.Submit("保存"), sg.Cancel("終了")],
#]
#表示バージョン（翻訳・保存・表示）
# layout = [
#    [sg.Text("英語で書かれているwordの文書を選択してください")],
#    [sg.Text("ファイル"), sg.InputText(), sg.FileBrowse(key="file1",button_text="ファイル選択")],
#    [sg.Submit("翻訳"), sg.Submit("保存"), sg.Submit("表示"), sg.Cancel("終了")],
# ]
#最終バージョン（翻訳・保存・表示・印刷）
#layout = [
#   [sg.Text("英語で書かれているwordの文書を選択してください")],
#   [sg.Text("ファイル"), sg.InputText(), sg.FileBrowse(key="file1",button_text="ファイル選択")],
#   [sg.Submit("翻訳"), sg.Submit("保存"), sg.Submit("表示"), sg.Submit("印刷"), sg.Cancel("終了")],
#]
#ファイル選択用画面の表示
window = sg.Window("ファイル選択", layout,disable_close=True)

while True:  # Event Loop
    event, values = window.read()
    if event == "終了":
        break
    #翻訳機能
    if event == "翻訳":
        #docxを開く
        doc = docx.Document(values['file1'])
        #新しいドキュメントオブジェクトを作成する
        newdoc = docx.Document()
        #段落を読み込む
        for paragraph in doc.paragraphs:
            #段落の文章をコンソールに出力する
            print(paragraph.text)
            if paragraph.text != "":
                #段落が空でなければ翻訳する
                response = translate.translate_text(Text=paragraph.text,
                        SourceLanguageCode=SRC_LANG, TargetLanguageCode=TRG_LANG )
                #変換結果をコンソールに出力する
                print(response['TranslatedText'])
                #誤変換をリプレイスする（本当は辞書登録が望ましい）
                paragraph.text = response['TranslatedText'].replace('証人', '本契約') \
                    .replace('一方、', '').replace('扇州','扇子')
            
            #翻訳後の文を段落に追加する
            newdoc.add_paragraph(paragraph.text)
            #文字の配置を元の文章に合わせる
            newdoc.paragraphs[-1].alignment = paragraph.alignment
    #wordへの保存機能
    if event == "保存" and 'newdoc' in locals():
        #wordファイルの保存を行う　コメント①
        newdoc.save(cwd + "翻訳結果.docx")
    #wordを起動して表示する
    if event == "表示" and os.path.isfile(cwd + "翻訳結果.docx"):
        #作成したwordファイルを表示　コメント②
        #Wordを起動する : Applicationオブジェクトを生成する
        Application=win32com.client.Dispatch("Word.Application")
        Application.Documents.Open(cwd + "翻訳結果.docx")
        Application.Visible=True
    #wordファイルを印刷する
    if event == "印刷":
        try:
            #pdfへ変換する　コメント③
            convert(cwd + "翻訳結果.docx")

            #acrobat.exe 保管場所　コメント③
            acr_path = "C:/Program Files/Adobe/Acrobat DC/Acrobat/Acrobat.exe"

            #acrobatで変換後のPDFを表示する　コメント③
            pdf_pro = subprocess.Popen([acr_path,cwd + "翻訳結果.pdf"], shell=False)
        except Exception:
            continue

#acrobatを終了する　コメント③
# pdf_pro.kill()